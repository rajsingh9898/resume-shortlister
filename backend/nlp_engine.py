import io
import re
import os
import json

_compiled_regex_cache = {}

def get_compiled_regex(pattern: str):
    global _compiled_regex_cache
    regex = _compiled_regex_cache.get(pattern)
    if regex is None:
        regex = re.compile(pattern)
        _compiled_regex_cache[pattern] = regex
    return regex


try:
    from backend.logger import logger
except ImportError:
    from logger import logger
try:
    import pypdf
except ImportError:
    import PyPDF2 as pypdf
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Hardcoded common English stopwords to ensure offline reliability
STOPWORDS = set([
    "i", "me", "my", "myself", "we", "our", "ours", "ourselves", "you", "your", "yours", "yourself", "yourselves",
    "he", "him", "his", "himself", "she", "her", "hers", "herself", "it", "its", "itself", "they", "them",
    "their", "theirs", "themselves", "what", "which", "who", "whom", "this", "that", "these", "those",
    "am", "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "having", "do", "does",
    "did", "doing", "a", "an", "the", "and", "but", "if", "or", "because", "as", "until", "while", "of",
    "at", "by", "for", "with", "about", "against", "between", "into", "through", "during", "before", "after",
    "above", "below", "to", "from", "up", "down", "in", "out", "on", "off", "over", "under", "again", "further",
    "then", "once", "here", "there", "when", "where", "why", "how", "all", "any", "both", "each", "few", "more",
    "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than", "too", "very",
    "s", "t", "can", "will", "just", "don", "should", "now", "of", "in", "to", "for", "with", "on", "at", "by",
    "from", "up", "about", "into", "over", "after"
])

# Dynamic loaded paths
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
TAXONOMY_PATH = os.path.join(CURRENT_DIR, "skills_taxonomy.json")

SKILLS_DB = {}
SKILL_SYNONYMS = {}

def load_skills_taxonomy():
    global SKILLS_DB, SKILL_SYNONYMS
    if os.path.exists(TAXONOMY_PATH):
        try:
            with open(TAXONOMY_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
                SKILLS_DB = data.get("skills_db", {})
                SKILL_SYNONYMS = data.get("skill_synonyms", {})
                return
        except Exception as e:
            logger.warning(f"Failed to load skills taxonomy: {e}. Using fallback defaults.")
            
    # Hardcoded default fallback
    SKILLS_DB = {
        "Languages": [
            "python", "javascript", "typescript", "java", "c\\+\\+", "c#", "ruby", "golang", "rust", 
            "php", "html", "css", "sql", "r", "swift", "kotlin", "scala", "perl", "bash", "shell"
        ],
        "Frameworks & Libraries": [
            "react", "angular", "vue", "next\\.js", "node\\.js", "express", "django", "flask", "fastapi", 
            "spring boot", "laravel", "rails", "asp\\.net", "tensorflow", "pytorch", "keras", "pandas", 
            "numpy", "scikit-learn", "scipy", "jquery", "bootstrap", "tailwind", "nextjs", "nodejs", 
            "spring", "dotnet", "react native", "flutter", "vuejs"
        ],
        "Databases & Tools": [
            "postgresql", "mysql", "mongodb", "redis", "sqlite", "oracle", "sql server", "dynamodb", 
            "elasticsearch", "cassandra", "firebase", "neo4j", "mariadb", "postgres"
        ],
        "Cloud & DevOps": [
            "aws", "azure", "gcp", "docker", "kubernetes", "git", "github", "gitlab", "jenkins", 
            "terraform", "ansible", "ci/cd", "linux", "nginx", "apache", "circleci", 
            "amazon web services", "google cloud"
        ],
        "Methodologies & Domains": [
            "agile", "scrum", "project management", "machine learning", "deep learning", "nlp", 
            "computer vision", "data analysis", "data science", "devops", "qa testing", "ui/ux", 
            "frontend", "backend", "full stack", "web development", "software engineering", "microservices",
            "rest api", "graphql", "system design", "artificial intelligence", "ai"
        ],
        "Soft Skills": [
            "communication", "leadership", "teamwork", "problem solving", "critical thinking", 
            "time management", "collaboration", "creativity", "presentation", "negotiation"
        ]
    }
    SKILL_SYNONYMS = {
        "PostgreSQL": ["postgresql", "postgres", "sql database", "psql"],
        "FastAPI": ["fastapi", "fast api", "asgi", "python asgi"],
        "Docker": ["docker", "containers", "containerization", "dockerfiles", "dockerize"],
        "Kubernetes": ["kubernetes", "k8s", "helm", "orchestration", "argocd"],
        "React": ["react", "reactjs", "react.js", "react-router", "redux"],
        "CI/CD": ["ci/cd", "pipeline", "pipelines", "jenkins", "github actions", "gitlab ci", "continuous integration"],
        "Git": ["git", "github", "gitlab", "bitbucket", "version control"],
        "MongoDB": ["mongodb", "mongo", "nosql", "document database"],
        "Python": ["python", "django", "flask", "fastapi", "asyncio"],
        "JavaScript": ["javascript", "js", "typescript", "ts", "es6"]
    }

LOWER_SYNONYMS_MAP = {}

def build_synonyms_map():
    global LOWER_SYNONYMS_MAP
    LOWER_SYNONYMS_MAP = {}
    for canonical_name, aliases in SKILL_SYNONYMS.items():
        all_names = [canonical_name.lower()] + [a.lower() for a in aliases]
        for name in all_names:
            if name in LOWER_SYNONYMS_MAP:
                LOWER_SYNONYMS_MAP[name].update(all_names)
            else:
                LOWER_SYNONYMS_MAP[name] = set(all_names)

load_skills_taxonomy()
build_synonyms_map()

# Lazy loading of sentence-transformers
_transformer_model = None

def get_transformer_model():
    global _transformer_model
    if _transformer_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            os.environ["TOKENIZERS_PARALLELISM"] = "false"
            _transformer_model = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as e:
            logger.warning(f"sentence-transformers load failed: {e}. Falling back to TF-IDF.")
            _transformer_model = False
    return _transformer_model

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        for page in reader.pages:
            val = page.extract_text()
            if val:
                text += val + "\n"
    except Exception as e:
        text = f"Error parsing PDF: {str(e)}"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        return f"Error parsing text file: {str(e)}"

def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    else:
        return extract_text_from_txt(file_bytes)

def preprocess_text(text: str) -> str:
    text = text.lower()
    words = get_compiled_regex(r'\b[a-z0-9#\+\-\.]+\b').findall(text)
    cleaned = [w for w in words if w not in STOPWORDS and len(w) > 1]
    return " ".join(cleaned)

def parse_experience_years_with_confidence(text: str) -> tuple:
    text_lower = text.lower()
    
    high_vals = []
    low_vals = []
    
    patterns_high = [
        r'(\d+(?:\.\d+)?)\s*(?:\+|plus)?\s*(?:years?|yrs?)\b(?:\s*(?:of)?\s*(?:experience|exp|work|industry|professional|in))\b',
        r'\b(?:experience|exp|work)\b\s*(?:of)?\s*(?:at\s+least|over)?\s*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\b',
        r'\btotal\b\s*(?:of)?\s*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\b'
    ]
    patterns_low = [
        r'\b(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\b'
    ]
    
    for pattern in patterns_high:
        matches = get_compiled_regex(pattern).findall(text_lower)
        for m in matches:
            try:
                val = float(m)
                if val < 40:
                    high_vals.append(val)
            except ValueError:
                pass
                
    for pattern in patterns_low:
        matches = get_compiled_regex(pattern).findall(text_lower)
        for m in matches:
            try:
                val = float(m)
                if val < 40:
                    low_vals.append(val)
            except ValueError:
                pass
                
    max_high = max(high_vals) if high_vals else 0.0
    max_low = max(low_vals) if low_vals else 0.0
    
    max_years = max(max_high, max_low)
    
    if max_years == 0.0:
        confidence = 0.95
    elif max_years == max_high:
        confidence = 0.9
    else:
        confidence = 0.5
        
    return max_years, confidence

def parse_experience_years(text: str) -> float:
    years, _ = parse_experience_years_with_confidence(text)
    return years

def parse_education_degrees_with_confidence(text: str) -> tuple:
    text_lower = text.lower()
    degrees = []
    
    degree_map_high = {
        "PhD": ["ph.d", "doctor of philosophy", "doctorate"],
        "Master": ["master of science", "master of arts", "master of business", "m.s.", "m.tech", "mba", "mca", "m.sc"],
        "Bachelor": ["bachelor of science", "bachelor of technology", "b.s.", "b.tech", "btech", "b.a.", "ba", "bca", "b.sc"]
    }
    
    degree_map_low = {
        "PhD": ["phd"],
        "Master": ["master", "ms", "msc", "mtech"],
        "Bachelor": ["bachelor", "bs", "bsc"]
    }
    
    confidence = 0.5
    matched_high = False
    
    for deg_type, terms in degree_map_high.items():
        for term in terms:
            pattern = r'\b' + re.escape(term) + r'\b'
            if term.endswith('.'):
                pattern = r'\b' + re.escape(term)
            if get_compiled_regex(pattern).search(text_lower):
                if deg_type not in degrees:
                    degrees.append(deg_type)
                    matched_high = True
                break
                
    for deg_type, terms in degree_map_low.items():
        if deg_type in degrees:
            continue
        for term in terms:
            pattern = r'\b' + re.escape(term) + r'\b'
            if get_compiled_regex(pattern).search(text_lower):
                if deg_type not in degrees:
                    degrees.append(deg_type)
                break
                
    if matched_high:
        confidence = 0.9
    elif degrees:
        confidence = 0.6
    else:
        confidence = 0.95
        
    return degrees, confidence
 
def parse_education_degrees(text: str) -> list:
    degrees, _ = parse_education_degrees_with_confidence(text)
    return degrees
 
def parse_soft_traits(text: str) -> list:
    text_lower = text.lower()
    traits = []
    trait_patterns = {
        "Leadership & Mentorship": [r"\bmanaged\b", r"\blead\b", r"\bmentor\b", r"\bspearheaded\b", r"\bdirected\b", r"\bleadership\b"],
        "System Design & Architecture": [r"\bscalability\b", r"\barchitecture\b", r"\bmicroservices\b", r"\bsystem design\b", r"\brefactor\b"],
        "Agile Delivery & DevOps": [r"\bagile\b", r"\bscrum\b", r"\bsprint\b", r"\bjira\b", r"\bdevops\b", r"\bci/cd\b"]
    }
    for trait, patterns in trait_patterns.items():
        for pat in patterns:
            if get_compiled_regex(pat).search(text_lower):
                traits.append(trait)
                break
    return traits
 
def extract_skills_from_text(text: str) -> dict:
    text_lower = text.lower()
    extracted = {}
    
    for category, skills in SKILLS_DB.items():
        matched = []
        for skill in skills:
            if '+' in skill or '#' in skill or '.' in skill:
                pattern = r'(?:^|\s|[.,/():\-])' + skill + r'(?:$|\s|[.,/():\-])'
            else:
                pattern = r'\b' + skill + r'\b'
                
            if get_compiled_regex(pattern).search(text_lower):
                clean_name = skill.replace("\\", "")
                
                if clean_name == "nextjs":
                    clean_name = "next.js"
                elif clean_name == "nodejs":
                    clean_name = "node.js"
                elif clean_name == "vuejs":
                    clean_name = "vue"
                elif clean_name == "postgres":
                    clean_name = "postgresql"
                elif clean_name == "dotnet":
                    clean_name = "asp.net"
                elif clean_name == "amazon web services":
                    clean_name = "aws"
                elif clean_name == "google cloud":
                    clean_name = "gcp"
                
                proper_cases = {
                    "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
                    "java": "Java", "c++": "C++", "c#": "C#", "ruby": "Ruby", "golang": "Go",
                    "rust": "Rust", "php": "PHP", "html": "HTML", "css": "CSS", "sql": "SQL",
                    "r": "R", "swift": "Swift", "kotlin": "Kotlin", "scala": "Scala",
                    "perl": "Perl", "bash": "Bash", "shell": "Shell", "react": "React",
                    "angular": "Angular", "vue": "Vue", "next.js": "Next.js", "node.js": "Node.js",
                    "express": "Express", "django": "Django", "flask": "Flask", "fastapi": "FastAPI",
                    "spring boot": "Spring Boot", "laravel": "Laravel", "rails": "Ruby on Rails",
                    "asp.net": "ASP.NET", "tensorflow": "TensorFlow", "pytorch": "PyTorch",
                    "keras": "Keras", "pandas": "Pandas", "numpy": "NumPy", "scikit-learn": "Scikit-Learn",
                    "scipy": "SciPy", "jquery": "jQuery", "bootstrap": "Bootstrap", "tailwind": "Tailwind",
                    "postgresql": "PostgreSQL", "mysql": "MySQL", "mongodb": "MongoDB", "redis": "Redis",
                    "sqlite": "SQLite", "oracle": "Oracle", "sql server": "SQL Server", "dynamodb": "DynamoDB",
                    "elasticsearch": "Elasticsearch", "cassandra": "Cassandra", "firebase": "Firebase",
                    "neo4j": "Neo4j", "mariadb": "MariaDB", "aws": "AWS", "azure": "Azure", "gcp": "GCP",
                    "docker": "Docker", "kubernetes": "Kubernetes", "git": "Git", "github": "GitHub",
                    "gitlab": "GitLab", "jenkins": "Jenkins", "terraform": "Terraform", "ansible": "Ansible",
                    "ci/cd": "CI/CD", "linux": "Linux", "nginx": "Nginx", "apache": "Apache",
                    "circleci": "CircleCI", "agile": "Agile", "scrum": "Scrum",
                    "project management": "Project Management", "machine learning": "Machine Learning",
                    "deep learning": "Deep Learning", "nlp": "NLP", "computer vision": "Computer Vision",
                    "data analysis": "Data Analysis", "data science": "Data Science", "devops": "DevOps",
                    "qa testing": "QA & Testing", "ui/ux": "UI/UX", "frontend": "Frontend",
                    "backend": "Backend", "full stack": "Full Stack", "web development": "Web Development",
                    "software engineering": "Software Engineering", "microservices": "Microservices",
                    "rest api": "REST APIs", "graphql": "GraphQL", "system design": "System Design",
                    "artificial intelligence": "AI", "communication": "Communication",
                    "leadership": "Leadership", "teamwork": "Teamwork", "problem solving": "Problem Solving",
                    "critical thinking": "Critical Thinking", "time management": "Time Management",
                    "collaboration": "Collaboration", "creativity": "Creativity",
                    "presentation": "Presentation", "negotiation": "Negotiation"
                }
                
                final_name = proper_cases.get(clean_name, clean_name.title())
                if final_name not in matched:
                    matched.append(final_name)
                    
        if matched:
            extracted[category] = sorted(matched)
            
    return extracted

def check_skill_match_raw(jd_skill: str, candidate_text: str, candidate_skills: set) -> bool:
    jd_lower = jd_skill.lower()
    cand_lower = {c.lower() for c in candidate_skills}
    
    # Get expanded synonyms list
    expanded_aliases = LOWER_SYNONYMS_MAP.get(jd_lower, {jd_lower})
    
    # 1. Check if candidate has direct skill or any synonym in parsed candidate_skills
    for alias in expanded_aliases:
        if alias in cand_lower:
            return True
            
    # 2. Check if candidate has direct skill or any synonym in raw text
    text_lower = candidate_text.lower()
    for alias in expanded_aliases:
        if '+' in alias or '#' in alias or '.' in alias:
            pattern = r'(?:^|\s|[.,/():\-])' + re.escape(alias) + r'(?:$|\s|[.,/():\-])'
        else:
            pattern = r'\b' + re.escape(alias) + r'\b'
            
        if get_compiled_regex(pattern).search(text_lower):
            return True
            
    return False

def pearson_correlation(x: list, y: list) -> float:
    n = len(x)
    if n < 2:
        return 0.0
    mean_x = sum(x) / n
    mean_y = sum(y) / n
    num = sum((x[i] - mean_x) * (y[i] - mean_y) for i in range(n))
    den_x = sum((x[i] - mean_x) ** 2 for i in range(n))
    den_y = sum((y[i] - mean_y) ** 2 for i in range(n))
    if den_x == 0 or den_y == 0:
        return 0.0
    return num / ((den_x * den_y) ** 0.5)

def compute_nlp_shortlist(jd_raw: str, resumes: list, semantic_weight: float = 0.5, team_profile: dict = None) -> dict:
    # 1. Parse Job Description Parameters
    jd_clean = preprocess_text(jd_raw)
    jd_skills_dict = extract_skills_from_text(jd_raw)
    jd_exp = parse_experience_years(jd_raw)
    jd_degrees = parse_education_degrees(jd_raw)
    
    jd_skills = []
    for cat_skills in jd_skills_dict.values():
        jd_skills.extend(cat_skills)
    jd_skills_set = set(jd_skills)
    
    if not jd_clean.strip():
        jd_clean = jd_raw.lower()
        
    cleaned_resumes = []
    for res in resumes:
        cleaned_resumes.append(preprocess_text(res['raw_text']))
        
    # 2. Calculate TF-IDF & Cosine Similarity
    tfidf_similarities = [0.0] * len(resumes)
    if jd_clean.strip() and any(r.strip() for r in cleaned_resumes):
        try:
            documents = [jd_clean] + cleaned_resumes
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform(documents)
            sim_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
            tfidf_similarities = sim_scores[0].tolist()
        except Exception:
            pass
            
    # Calculate Sentence-Transformers embeddings cosine similarity
    semantic_similarities = [0.0] * len(resumes)
    model = get_transformer_model()
    if model:
        try:
            import torch
            from sentence_transformers.util import cos_sim
            jd_emb = model.encode(jd_raw, convert_to_tensor=True)
            res_embs = model.encode([r['raw_text'] for r in resumes], convert_to_tensor=True)
            sims = cos_sim(jd_emb, res_embs)
            semantic_similarities = sims[0].cpu().tolist()
        except Exception as e:
            logger.warning(f"Transformer similarity computations failed: {e}")
            
    # 3. Calculate candidate scores and compile analysis reports
    candidates_list = []
    for idx, res in enumerate(resumes):
        raw_txt = res['raw_text']
        c_skills_dict = extract_skills_from_text(raw_txt)
        
        c_skills = []
        for cat_skills in c_skills_dict.values():
            c_skills.extend(cat_skills)
        c_skills_set = set(c_skills)
        
        matched_skills = []
        missing_skills = []
        for req_skill in jd_skills_set:
            if check_skill_match_raw(req_skill, raw_txt, c_skills_set):
                matched_skills.append(req_skill)
            else:
                missing_skills.append(req_skill)
                
        matched_skills.sort()
        missing_skills.sort()
        
        skills_score = 0.0
        if jd_skills_set:
            skills_score = len(matched_skills) / len(jd_skills_set)
            
        candidate_exp, exp_conf = parse_experience_years_with_confidence(raw_txt)
        experience_score = 0.0
        if jd_exp > 0.0:
            if candidate_exp >= jd_exp:
                experience_score = 1.0
            else:
                experience_score = candidate_exp / jd_exp
        else:
            experience_score = 1.0
            
        candidate_degrees, deg_conf = parse_education_degrees_with_confidence(raw_txt)
        
        # Degree value hierarchy check for matches/exceeds logic
        degree_hierarchy = {"PhD": 3, "Master": 2, "Bachelor": 1}
        max_jd_deg_val = max([degree_hierarchy.get(d, 0) for d in jd_degrees]) if jd_degrees else 0
        max_cand_deg_val = max([degree_hierarchy.get(d, 0) for d in candidate_degrees]) if candidate_degrees else 0
        
        if max_jd_deg_val == 0:
            degree_match_score = 1.0
            degree_match = True
        elif max_cand_deg_val >= max_jd_deg_val:
            degree_match_score = 1.0
            degree_match = True
        elif max_cand_deg_val > 0:
            degree_match_score = 0.5
            degree_match = False
        else:
            degree_match_score = 0.0
            degree_match = False
            
        soft_traits = parse_soft_traits(raw_txt)
        soft_skills_score = len(soft_traits) / 3.0  # 3 categories max
        
        tfidf_sim = max(0.0, min(1.0, tfidf_similarities[idx]))
        semantic_sim = max(0.0, min(1.0, semantic_similarities[idx]))
        
        # Blend similarities: use sentence-transformers if successfully loaded
        if model:
            cosine_sim = (1.0 - semantic_weight) * tfidf_sim + semantic_weight * semantic_sim
        else:
            cosine_sim = tfidf_sim
            
        # Hybrid Scoring Blending (40% Semantic, 30% Keyword, 30% Rules)
        semantic_component = cosine_sim
        keyword_component = (tfidf_sim * 0.4) + (skills_score * 0.6)
        rule_component = (experience_score * 0.5) + (degree_match_score * 0.4) + (soft_skills_score * 0.1)
        
        final_score = (semantic_component * 0.4) + (keyword_component * 0.3) + (rule_component * 0.3)
        
        # 1. Domain Fit calculation
        domains = {
            "Finance": ["finance", "banking", "fintech", "ledger", "trade", "trading", "investment", "bank", "accountant", "accounting", "audit"],
            "Healthcare": ["healthcare", "medical", "clinical", "doctor", "patient", "health", "biotech", "pharma", "hospital", "nursing"],
            "E-commerce": ["e-commerce", "ecommerce", "retail", "shop", "sales", "checkout", "cart", "stripe", "payment", "order"],
            "Tech": ["software", "developer", "engineering", "web", "saas", "cloud", "infrastructure"]
        }
        jd_domains = []
        for dom_name, keywords in domains.items():
            if any(kw in jd_raw.lower() for kw in keywords):
                jd_domains.append(dom_name)
        domain_fit_score = 1.0
        if jd_domains:
            matched_domains = 0
            for dom_name in jd_domains:
                keywords = domains[dom_name]
                if any(kw in raw_txt.lower() for kw in keywords):
                    matched_domains += 1
            domain_fit_score = matched_domains / len(jd_domains) if jd_domains else 1.0
            domain_fit_score = max(0.6, domain_fit_score)
            
        # 2. Seniority Fit calculation
        seniority_fit_score = 1.0
        if jd_exp > 0.0:
            exp_diff = candidate_exp - jd_exp
            if exp_diff >= 0:
                seniority_fit_score = 1.0
            elif exp_diff >= -2.0:
                seniority_fit_score = 0.8
            else:
                seniority_fit_score = max(0.3, round(candidate_exp / jd_exp, 2))
        else:
            seniority_fit_score = 1.0
        is_jd_senior = any(x in jd_raw.lower() for x in ["senior", "lead", "architect", "principal", "manager"])
        is_cand_senior = any(x in raw_txt.lower() for x in ["senior", "lead", "architect", "principal", "manager"])
        if is_jd_senior and is_cand_senior:
            seniority_fit_score = min(1.0, seniority_fit_score + 0.1)
            
        # 3. Culture/Soft-Signal Fit
        culture_fit_score = 0.5
        if len(soft_traits) >= 2:
            culture_fit_score = 1.0
        elif len(soft_traits) == 1:
            culture_fit_score = 0.75
        else:
            soft_skills_terms = ["communication", "teamwork", "adaptability", "problem solving", "collaboration", "motivated"]
            matched_terms = [t for t in soft_skills_terms if t in raw_txt.lower()]
            culture_fit_score = min(0.9, 0.4 + 0.1 * len(matched_terms))

        # 4. Job-to-Team Fit calculation
        team_fit_score = 1.0
        team_fit_details = {
            "mindset_alignment": "Startup (Match)",
            "focus_alignment": "Backend-heavy (Match)",
            "expectation_alignment": "Ownership (Match)"
        }
        if team_profile:
            # Mindset
            mindset = team_profile.get("mindset", "Enterprise")
            mindset_score = 0.5
            if mindset == "Startup":
                startup_keywords = ["startup", "prototype", "mvp", "rapid", "agile", "ownership", "wear multiple hats", "fast-paced", "built from scratch"]
                matched_count = sum(1 for kw in startup_keywords if kw in raw_txt.lower())
                startup_tech = ["react", "node", "python", "fastapi", "django", "typescript", "nextjs", "vue"]
                tech_count = sum(1 for t in startup_tech if t in raw_txt.lower())
                if matched_count > 0 or tech_count >= 2:
                    mindset_score = 1.0
                    team_fit_details["mindset_alignment"] = "Startup (Match)"
                else:
                    mindset_score = 0.6
                    team_fit_details["mindset_alignment"] = "Startup (Low Correlation)"
            else:
                ent_keywords = ["enterprise", "architecture", "robust", "scalable", "compliance", "process", "legacy", "migration", "testing", "unit test", "java", "c#", "dotnet", "oracle", "kubernetes"]
                matched_count = sum(1 for kw in ent_keywords if kw in raw_txt.lower())
                if matched_count >= 2:
                    mindset_score = 1.0
                    team_fit_details["mindset_alignment"] = "Enterprise (Match)"
                else:
                    mindset_score = 0.7
                    team_fit_details["mindset_alignment"] = "Enterprise (Partial Match)"
            
            # Focus
            focus = team_profile.get("focus", "Backend-heavy")
            focus_score = 0.5
            backend_terms = ["python", "java", "c#", "postgres", "sql", "django", "fastapi", "database", "api", "backend", "node", "docker", "redis", "celery"]
            frontend_terms = ["react", "javascript", "html", "css", "frontend", "ui", "ux", "user interface", "design", "tailwind", "sass"]
            be_count = sum(1 for t in backend_terms if t in raw_txt.lower())
            fe_count = sum(1 for t in frontend_terms if t in raw_txt.lower())
            if focus == "Backend-heavy":
                if be_count > fe_count:
                    focus_score = 1.0
                    team_fit_details["focus_alignment"] = "Backend-heavy (Match)"
                elif be_count > 0:
                    focus_score = 0.8
                    team_fit_details["focus_alignment"] = "Backend-heavy (Partial Match)"
                else:
                    focus_score = 0.4
                    team_fit_details["focus_alignment"] = "Backend-heavy (Low Alignment)"
            elif focus == "Product-heavy":
                if fe_count > be_count:
                    focus_score = 1.0
                    team_fit_details["focus_alignment"] = "Product-heavy (Match)"
                elif fe_count > 0:
                    focus_score = 0.8
                    team_fit_details["focus_alignment"] = "Product-heavy (Partial Match)"
                else:
                    focus_score = 0.4
                    team_fit_details["focus_alignment"] = "Product-heavy (Low Alignment)"
            else:
                if be_count > 0 and fe_count > 0:
                    focus_score = 1.0
                    team_fit_details["focus_alignment"] = "Fullstack (Match)"
                elif be_count > 0 or fe_count > 0:
                    focus_score = 0.7
                    team_fit_details["focus_alignment"] = "Fullstack (Partial Match)"
                else:
                    focus_score = 0.4
                    team_fit_details["focus_alignment"] = "Fullstack (Low Alignment)"
                    
            # Expectation
            expectation = team_profile.get("expectation", "Ownership")
            expectation_score = 0.5
            if expectation == "Ownership":
                own_keywords = ["lead", "mentor", "manager", "owner", "spearheaded", "designed", "architected", "leadership", "delivery", "initiative"]
                matched_count = sum(1 for kw in own_keywords if kw in raw_txt.lower())
                if matched_count > 0:
                    expectation_score = 1.0
                    team_fit_details["expectation_alignment"] = "Ownership (Match)"
                else:
                    expectation_score = 0.5
                    team_fit_details["expectation_alignment"] = "Ownership (Partial)"
            else:
                sup_keywords = ["support", "maintenance", "ticket", "resolve", "debug", "document", "jira", "customer support", "fixing", "bugs"]
                matched_count = sum(1 for kw in sup_keywords if kw in raw_txt.lower())
                if matched_count > 0:
                    expectation_score = 1.0
                    team_fit_details["expectation_alignment"] = "Support (Match)"
                else:
                    expectation_score = 0.6
                    team_fit_details["expectation_alignment"] = "Support (Partial)"
                    
            team_fit_score = round((mindset_score + focus_score + expectation_score) / 3.0, 2)
            
            # Blend team fit score (20% weight) into the final ranking score
            final_score = (final_score * 0.8) + (team_fit_score * 0.2)
            
        # Generate Explainable AI "Why this candidate?" statement
        matched_str = ", ".join(matched_skills[:2]) if matched_skills else "none"
        why_candidate = ""
        if final_score > 0.75:
            why_candidate = f"Highly Recommended: This candidate exhibits excellent skills alignment matching {matched_str}. They have a strong experience match of {candidate_exp:.1f} years. "
        elif final_score > 0.45:
            why_candidate = f"Good Match: This candidate fits the core requirements with {candidate_exp:.1f} years of experience. They cover skills like {matched_str}. "
        else:
            why_candidate = f"Underqualified: The candidate has low skill coverage, missing key technologies. Experience is {candidate_exp:.1f} years. "
            
        if team_profile:
            why_candidate += f"Additionally, they show high compatibility with the team's {team_profile.get('focus', 'Backend-heavy')} focus, demonstrating a {team_profile.get('mindset', 'Enterprise')} mindset."
        else:
            why_candidate += "No team profile alignment details were configured."
        
        # Generate score explainability fields (positive/negative indicators)
        reasons_high = []
        reasons_low = []
        
        if semantic_component > 0.6:
            reasons_high.append("Excellent semantic alignment with the job description context.")
        elif semantic_component < 0.25:
            reasons_low.append("Low contextual alignment with the overall job scope.")
            
        if matched_skills:
            top_matched = ", ".join(matched_skills[:3])
            reasons_high.append(f"Strong match for required skills: {top_matched}.")
        if missing_skills:
            top_missing = ", ".join(missing_skills[:3])
            reasons_low.append(f"Missing required skills: {top_missing}.")
            
        if jd_exp > 0.0:
            if candidate_exp >= jd_exp:
                reasons_high.append(f"Exceeds/meets required experience: {candidate_exp:.1f} years matched (JD asked: {jd_exp:.1f} years).")
            else:
                reasons_low.append(f"Experience ({candidate_exp:.1f} years) is lower than required {jd_exp:.1f} years.")
                
        if jd_degrees:
            if degree_match_score == 1.0:
                reasons_high.append("Academic qualification matches or exceeds the required degree level.")
            elif degree_match_score == 0.5:
                reasons_low.append("Academic qualification level is lower than required degree level.")
            else:
                reasons_low.append("Missing required academic degree qualifications.")
                
        if len(soft_traits) >= 2:
            reasons_high.append("Demonstrates key traits: leadership, system architecture, or agile experience.")
            
        explainability = {
            "reasons_high": reasons_high,
            "reasons_low": reasons_low,
            "breakdown": {
                "skills": round(skills_score * 100, 1),
                "experience": round(experience_score * 100, 1),
                "domain_fit": round(domain_fit_score * 100, 1),
                "seniority_fit": round(seniority_fit_score * 100, 1),
                "soft_signals": round(culture_fit_score * 100, 1),
                "team_fit": round(team_fit_score * 100, 1)
            },
            "team_fit_details": team_fit_details,
            "why_candidate": why_candidate
        }
        
        candidates_list.append({
            "filename": res['filename'],
            "score": round(final_score * 100, 1),
            "cosine_score": round(cosine_sim * 100, 1),
            "skills_score": round(skills_score * 100, 1),
            "experience_score": round(experience_score * 100, 1),
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "all_extracted_skills": c_skills_dict,
            "candidate_exp": candidate_exp,
            "experience_confidence": round(exp_conf, 2),
            "candidate_degrees": candidate_degrees,
            "degrees_confidence": round(deg_conf, 2),
            "degree_match": degree_match,
            "soft_traits": soft_traits,
            "model_version": "v2.1.0",
            "explainability": explainability,
            "snippet": raw_txt[:400] + ("..." if len(raw_txt) > 400 else "")
        })
        
    candidates_list.sort(key=lambda x: x['score'], reverse=True)
    
    # 4. Post-Scoring Bias & Fairness Checker
    bias_warnings = []
    if len(candidates_list) >= 2:
        scores = [c['score'] for c in candidates_list]
        lengths = [len(r['raw_text']) for r in resumes]
        gaps = [1.0 if any(g in r['raw_text'].lower() for g in ["career break", "career gap", "employment gap", "sabbatical", "parental leave"]) else 0.0 for r in resumes]
        
        # Calculate formatting flags (special char ratio or very short length)
        formatting_flags = []
        for r in resumes:
            txt = r['raw_text']
            special_count = len(get_compiled_regex(r'[^a-zA-Z0-9\s]').findall(txt))
            total_count = len(txt) if len(txt) > 0 else 1
            ratio = special_count / total_count
            if ratio > 0.15 or len(txt) < 200:
                formatting_flags.append(1.0)
            else:
                formatting_flags.append(0.0)
                
        # Length correlation check
        corr_len = pearson_correlation(scores, lengths)
        if abs(corr_len) > 0.5:
            bias_warnings.append(f"⚠️ Bias Alert: Match scores correlate strongly with resume length (correlation: {corr_len:.2f}). Longer resumes may have an unfair advantage.")
            
        # Career breaks check
        corr_gaps = pearson_correlation(scores, gaps)
        if corr_gaps < -0.4:
            bias_warnings.append(f"⚠️ Bias Alert: Candidate scores are negatively correlated with career breaks or employment gaps (correlation: {corr_gaps:.2f}). System may be penalizing gaps.")
            
        # Formatting check
        corr_format = pearson_correlation(scores, formatting_flags)
        if corr_format < -0.4:
            bias_warnings.append(f"⚠️ Bias Alert: Match scores correlate negatively with non-standard formatting (correlation: {corr_format:.2f}). Formatting issues may be penalizing candidates.")
            
    return {
        "candidates": candidates_list,
        "bias_warnings": bias_warnings,
        "jd_requirements": {
            "skills": sorted(list(jd_skills_set)),
            "experience_years": jd_exp,
            "degrees": jd_degrees
        }
    }
